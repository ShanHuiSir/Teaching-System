import json
import re
from pathlib import Path
from typing import Optional, Tuple, Union

from docx import Document
from docx.shared import Pt

NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

_RE_ORDERED_DECIMAL = re.compile(r"^(\d+)[.、．)）]\s*")
_RE_ORDERED_PAREN = re.compile(r"^[（(]\d+[)）]\s*")
_RE_ORDERED_CHINESE = re.compile(r"^[一二三四五六七八九十]+[、．.]\s*")
_RE_UNORDERED = re.compile(r"^[-•·▪▸►○●◆◇]\s*")

_RE_STRIP_DECIMAL = re.compile(r"^\d+[.、．)）]\s*")
_RE_STRIP_CHINESE = re.compile(r"^[一二三四五六七八九十]+[、．.]\s*")
_RE_STRIP_PAREN = re.compile(r"^[（(]\d+[)）]\s*")


# ---------------------------------------------------------------------------
# Body element iteration
# ---------------------------------------------------------------------------

def _iter_body_elements(doc):
    body_children = list(doc.element.body)
    elements = []
    for child in body_children:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            for p in doc.paragraphs:
                if p._element is child:
                    elements.append(("paragraph", p))
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._element is child:
                    elements.append(("table", t))
                    break
    return elements


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------

def _get_paragraph(elements, idx):
    if idx >= len(elements):
        return None, None
    el_type, el = elements[idx]
    return (el, el_type) if el_type == "paragraph" else (None, el_type)


def _clean_item_text(text: str, ordered: bool) -> str:
    if not ordered:
        return text
    for pat in (_RE_STRIP_DECIMAL, _RE_STRIP_CHINESE, _RE_STRIP_PAREN):
        cleaned = pat.sub("", text, count=1)
        if cleaned != text:
            return cleaned
    return text


def _detect_level(paragraph) -> Optional[int]:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading "):
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return None
    if style_name == "Normal":
        for run in paragraph.runs:
            if run.bold and run.font.size and run.font.size >= Pt(20):
                return 1
    return None


# ---------------------------------------------------------------------------
# Word native list detection
# ---------------------------------------------------------------------------

def _get_word_list_info(paragraph) -> Tuple[Optional[str], int]:
    pPr = paragraph._element.find(f"{NS}pPr")
    if pPr is None:
        return None, 0
    numPr = pPr.find(f"{NS}numPr")
    if numPr is None:
        return None, 0
    numId_el = numPr.find(f"{NS}numId")
    if numId_el is None:
        return None, 0
    numId = numId_el.get(f"{NS}val")
    ilvl_el = numPr.find(f"{NS}ilvl")
    ilvl = int(ilvl_el.get(f"{NS}val")) if ilvl_el is not None else 0
    return numId, ilvl


def _build_num_fmt_cache(doc) -> dict[str, str]:
    cache = {}
    try:
        numbering = doc.part.numbering_part._element
    except (KeyError, NotImplementedError):
        return cache
    for num in numbering.findall(f".//{NS}num"):
        num_id = num.get(f"{NS}numId")
        abstract_ref = num.find(f"{NS}abstractNumId")
        if abstract_ref is None:
            continue
        an_id = abstract_ref.get(f"{NS}val")
        for an in numbering.findall(f".//{NS}abstractNum"):
            if an.get(f"{NS}abstractNumId") == an_id:
                lvl0 = an.find(f"{NS}lvl[@{{{NS}}}ilvl='0']")
                if lvl0 is None:
                    lvl0 = an.find(f"{NS}lvl")
                if lvl0 is not None:
                    num_fmt = lvl0.find(f"{NS}numFmt")
                    if num_fmt is not None:
                        cache[num_id] = num_fmt.get(f"{NS}val")
                break
    return cache


def _list_is_ordered(num_fmt_cache: dict[str, str], num_id: str) -> bool:
    fmt = num_fmt_cache.get(num_id)
    if fmt is None:
        return True
    return fmt not in ("bullet",)


# ---------------------------------------------------------------------------
# Manual list detection
# ---------------------------------------------------------------------------

def _detect_manual_list(text: str) -> Optional[Tuple[str, str]]:
    if _RE_ORDERED_DECIMAL.match(text):
        return "ordered:decimal", _RE_ORDERED_DECIMAL.sub("", text, count=1)
    if _RE_ORDERED_PAREN.match(text):
        return "ordered:paren", _RE_ORDERED_PAREN.sub("", text, count=1)
    if _RE_ORDERED_CHINESE.match(text):
        return "ordered:chinese", _RE_ORDERED_CHINESE.sub("", text, count=1)
    if _RE_UNORDERED.match(text):
        return "unordered:bullet", _RE_UNORDERED.sub("", text, count=1)
    return None


# ---------------------------------------------------------------------------
# Table extraction
# ---------------------------------------------------------------------------

def _table_to_content(table) -> dict:
    rows = table.rows
    headers = [cell.text.strip() for cell in rows[0].cells] if rows else []
    data = [[cell.text.strip() for cell in row.cells] for row in rows[1:]] if len(rows) > 1 else []
    return {"type": "table", "headers": headers, "rows": data}


# ---------------------------------------------------------------------------
# List consumers
# ---------------------------------------------------------------------------

def _consume_word_list(elements, start_idx, num_fmt_cache, first_num_id, first_ilvl):
    items = []
    idx = start_idx
    base_num_id = first_num_id

    while idx < len(elements):
        p, el_type = _get_paragraph(elements, idx)
        if p is None:
            break
        text = p.text.strip()
        if not text:
            idx += 1
            continue

        num_id, ilvl = _get_word_list_info(p)
        if num_id is None:
            break
        if num_id != base_num_id and ilvl == 0 and items:
            break
        items.append((ilvl, text, num_id))
        idx += 1

    if not items:
        return None, idx

    ordered = _list_is_ordered(num_fmt_cache, base_num_id)
    result_items = _build_list_tree(items, ordered)
    return {"type": "list", "ordered": ordered, "items": result_items}, idx


def _build_list_tree(flat_items, ordered):
    result = []
    base_ilvl = flat_items[0][0]

    for ilvl, text, _num_id in flat_items:
        text = _clean_item_text(text, ordered)
        if ilvl == base_ilvl:
            result.append({"text": text})
        elif ilvl > base_ilvl:
            prev_item = result[-1]
            if "children" not in prev_item:
                prev_item["children"] = []
            prev_item["children"].append({"text": text})

    for item in result:
        if "children" in item:
            item["children"] = _group_items_into_lists(item["children"])

    return result


def _group_items_into_lists(raw_items):
    if not raw_items:
        return raw_items
    return [{"type": "list", "ordered": True, "items": raw_items}]


def _consume_manual_list(elements, start_idx, first_match):
    pattern_key, cleaned_text = first_match
    items = [{"text": cleaned_text}]
    idx = start_idx + 1

    while idx < len(elements):
        p, el_type = _get_paragraph(elements, idx)
        if p is None:
            break
        text = p.text.strip()
        if not text:
            idx += 1
            continue
        if _detect_level(p) is not None:
            break
        num_id, _ = _get_word_list_info(p)
        if num_id is not None:
            break
        match = _detect_manual_list(text)
        if match and match[0] == pattern_key:
            items.append({"text": match[1]})
            idx += 1
        else:
            break

    ordered = pattern_key.startswith("ordered")
    return {"type": "list", "ordered": ordered, "items": items}, idx


# ---------------------------------------------------------------------------
# Content consumption
# ---------------------------------------------------------------------------

def _consume_content(elements, start_idx, num_fmt_cache):
    content = []
    idx = start_idx

    while idx < len(elements):
        el_type, el = elements[idx]

        if el_type == "table":
            content.append(_table_to_content(el))
            idx += 1
            continue

        p = el
        text = p.text.strip()

        if not text:
            idx += 1
            continue

        if _detect_level(p) is not None:
            break

        num_id, ilvl = _get_word_list_info(p)
        if num_id is not None:
            list_item, idx = _consume_word_list(elements, idx, num_fmt_cache, num_id, ilvl)
            content.append(list_item)
            continue

        manual_match = _detect_manual_list(text)
        if manual_match:
            list_item, idx = _consume_manual_list(elements, idx, manual_match)
            content.append(list_item)
            continue

        content.append({"type": "paragraph", "text": text})
        idx += 1

    return content, idx


# ---------------------------------------------------------------------------
# Tree building
# ---------------------------------------------------------------------------

def _build_tree(doc):
    elements = _iter_body_elements(doc)
    num_fmt_cache = _build_num_fmt_cache(doc)
    root = []
    stack = []
    idx = 0

    while idx < len(elements):
        el_type, el = elements[idx]

        if el_type == "table":
            root.append(_table_to_content(el))
            idx += 1
            continue

        p = el
        text = p.text.strip()

        if not text:
            idx += 1
            continue

        level = _detect_level(p)

        if level is not None:
            node = {"level": level, "heading": text, "content": [], "children": []}
            while stack and stack[-1][0] >= level:
                stack.pop()
            if stack:
                stack[-1][1]["children"].append(node)
            else:
                root.append(node)
            stack.append((level, node))
            idx += 1
            content_items, idx = _consume_content(elements, idx, num_fmt_cache)
            node["content"] = content_items
        else:
            content_items, idx = _consume_content(elements, idx, num_fmt_cache)
            root.extend(content_items)

    return root


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def convert(filepath: Union[str, Path]) -> str:
    doc = Document(str(filepath))
    tree = _build_tree(doc)
    return json.dumps(tree, ensure_ascii=False, indent=2)


def convert_obj(filepath: Union[str, Path]):
    """Return Python object (list/dict) instead of JSON string."""
    doc = Document(str(filepath))
    return _build_tree(doc)
